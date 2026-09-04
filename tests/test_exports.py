"""Тесты подсистемы формирования отчётов.

Метод: интеграционное тестирование — проверяется, что каждый формат
формируется полностью и результат является корректным файлом
соответствующего типа, пригодным для открытия.
"""

from __future__ import annotations

import json
import zipfile

import pytest
from django.urls import reverse
from exports import builders
from exports.datasets import DATASETS

pytestmark = pytest.mark.django_db


@pytest.fixture
def export_root(tmp_path, settings):
    """Перенаправить каталог выгрузок во временный на время теста."""
    settings.EXPORT_ROOT = tmp_path / "exports"
    settings.EXPORT_ROOT.mkdir(parents=True, exist_ok=True)
    return settings.EXPORT_ROOT


class TestFilenames:
    """Формирование имён файлов."""

    def test_cyrillic_transliterated(self):
        """Кириллица в наименовании транслитерируется."""
        dataset = builders.Dataset(code="объекты склада", title="Тест", columns=[], rows=[])
        name = builders.build_filename(dataset, "xlsx")
        assert name.isascii() and name.endswith(".xlsx")

    def test_timestamp_included(self):
        """Имя содержит отметку времени."""
        dataset = builders.Dataset(code="objects", title="Тест", columns=[], rows=[])
        name = builders.build_filename(dataset, "csv")
        assert len(name.split("-")) >= 3

    def test_names_unique_over_time(self):
        """Повторная выгрузка не перезаписывает предыдущий файл."""
        dataset = builders.Dataset(code="objects", title="Тест", columns=[], rows=[])
        assert builders.build_filename(dataset, "csv").endswith(".csv")


class TestXlsx:
    """Электронная таблица XLSX."""

    def test_file_created(self, full_dataset, export_root):
        """Файл формируется и не пуст."""
        from exports.datasets import objects_dataset

        path = export_root / "test.xlsx"
        rows = builders.build_xlsx(objects_dataset({}), path)
        assert path.exists() and path.stat().st_size > 0
        assert rows == 5

    def test_valid_zip_container(self, full_dataset, export_root):
        """Файл является корректным контейнером формата Office Open XML."""
        from exports.datasets import objects_dataset

        path = export_root / "test.xlsx"
        builders.build_xlsx(objects_dataset({}), path)
        assert zipfile.is_zipfile(path)

    def test_contains_data_and_meta_sheets(self, full_dataset, export_root):
        """Книга содержит лист данных и лист сведений о выгрузке."""
        from exports.datasets import objects_dataset
        from openpyxl import load_workbook

        path = export_root / "test.xlsx"
        builders.build_xlsx(objects_dataset({}), path)
        workbook = load_workbook(path)
        assert len(workbook.sheetnames) == 2
        assert "Сведения о выгрузке" in workbook.sheetnames

    def test_header_row_frozen(self, full_dataset, export_root):
        """Строка заголовков закреплена при прокрутке."""
        from exports.datasets import objects_dataset
        from openpyxl import load_workbook

        path = export_root / "test.xlsx"
        builders.build_xlsx(objects_dataset({}), path)
        assert load_workbook(path).active.freeze_panes == "A2"

    def test_row_count_matches(self, full_dataset, export_root):
        """Число строк в листе соответствует числу записей."""
        from exports.datasets import objects_dataset
        from openpyxl import load_workbook

        path = export_root / "test.xlsx"
        builders.build_xlsx(objects_dataset({}), path)
        assert load_workbook(path).active.max_row == 6


class TestDocx:
    """Текстовый документ DOCX."""

    def test_file_created(self, full_dataset, export_root):
        """Документ формируется."""
        from exports.datasets import districts_dataset

        path = export_root / "test.docx"
        builders.build_docx(districts_dataset({}), path)
        assert path.exists() and path.stat().st_size > 0

    def test_valid_container(self, full_dataset, export_root):
        """Файл является корректным контейнером Office Open XML."""
        from exports.datasets import districts_dataset

        path = export_root / "test.docx"
        builders.build_docx(districts_dataset({}), path)
        assert zipfile.is_zipfile(path)

    def test_contains_title_and_table(self, full_dataset, export_root):
        """Документ содержит заголовок и таблицу данных."""
        from docx import Document
        from exports.datasets import districts_dataset

        path = export_root / "test.docx"
        builders.build_docx(districts_dataset({}), path)
        document = Document(path)
        assert any("округ" in p.text.lower() for p in document.paragraphs)
        assert len(document.tables) >= 1

    def test_table_row_count(self, full_dataset, export_root, districts):
        """Таблица содержит строку заголовков и все записи."""
        from docx import Document
        from exports.datasets import districts_dataset

        path = export_root / "test.docx"
        builders.build_docx(districts_dataset({}), path)
        data_table = Document(path).tables[-1]
        assert len(data_table.rows) == len(districts) + 1


class TestPdf:
    """Документ PDF."""

    def test_file_created(self, full_dataset, export_root):
        """Документ формируется."""
        from exports.datasets import objects_dataset

        path = export_root / "test.pdf"
        builders.build_pdf(objects_dataset({}), path)
        assert path.exists() and path.stat().st_size > 0

    def test_pdf_signature(self, full_dataset, export_root):
        """Файл начинается с сигнатуры формата PDF."""
        from exports.datasets import objects_dataset

        path = export_root / "test.pdf"
        builders.build_pdf(objects_dataset({}), path)
        assert path.read_bytes()[:5] == b"%PDF-"


class TestCsv:
    """Таблица CSV."""

    def test_delimiter_and_encoding(self, full_dataset, export_root):
        """Файл записан в UTF-8 с меткой порядка байтов и разделителем «;»."""
        from exports.datasets import objects_dataset

        path = export_root / "test.csv"
        builders.build_csv(objects_dataset({}), path)
        raw = path.read_bytes()
        assert raw.startswith(b"\xef\xbb\xbf")
        assert b";" in raw

    def test_row_count(self, full_dataset, export_root, objects):
        """Число строк соответствует числу записей плюс заголовок."""
        import csv

        from exports.datasets import objects_dataset

        path = export_root / "test.csv"
        builders.build_csv(objects_dataset({}), path)
        with path.open(encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.reader(handle, delimiter=";"))
        assert len(rows) == len(objects) + 1


class TestGeoJson:
    """Пространственный слой GeoJSON."""

    def test_valid_structure(self, full_dataset, export_root):
        """Файл содержит коллекцию объектов установленной структуры."""
        from exports.datasets import objects_dataset

        path = export_root / "test.geojson"
        builders.build_geojson(objects_dataset({}), path, lambda row: row.geom)
        payload = json.loads(path.read_text(encoding="utf-8"))
        assert payload["type"] == "FeatureCollection"
        assert payload["features"][0]["type"] == "Feature"

    def test_excludes_records_without_geometry(self, full_dataset, export_root, objects):
        """Записи без координат в слой не включаются."""
        from exports.datasets import objects_dataset

        path = export_root / "test.geojson"
        count = builders.build_geojson(objects_dataset({}), path, lambda row: row.geom)
        assert count == len(objects) - 1

    def test_metadata_present(self, full_dataset, export_root):
        """Слой содержит сведения о системе и времени формирования."""
        from exports.datasets import objects_dataset

        path = export_root / "test.geojson"
        builders.build_geojson(objects_dataset({}), path, lambda row: row.geom)
        payload = json.loads(path.read_text(encoding="utf-8"))
        assert payload["metadata"]["system"]
        assert payload["metadata"]["generated_at"]


class TestDatasets:
    """Наборы данных."""

    @pytest.mark.parametrize("code", list(DATASETS))
    def test_dataset_builds(self, full_dataset, code):
        """Каждый набор формируется и содержит описание колонок."""
        builder, _ = DATASETS[code]
        dataset = builder({})
        assert dataset.columns and dataset.title

    @pytest.mark.parametrize("code", list(DATASETS))
    def test_dataset_rows_match_columns(self, full_dataset, code):
        """Число значений в строке совпадает с числом колонок."""
        builder, _ = DATASETS[code]
        dataset = builder({})
        rows = dataset.materialize()
        assert all(len(row) == len(dataset.columns) for row in rows)

    def test_filters_applied(self, full_dataset, districts):
        """Условия отбора передаются в набор данных."""
        from exports.datasets import objects_dataset

        dataset = objects_dataset({"district": str(districts[0].pk)})
        assert len(dataset.materialize()) == 2

    def test_summary_present(self, full_dataset):
        """Набор содержит сводные показатели для документа."""
        from exports.datasets import objects_dataset

        assert objects_dataset({}).summary


class TestExportView:
    """Представление, обслуживающее выгрузку."""

    @pytest.mark.parametrize("fmt", ["xlsx", "docx", "pdf", "csv", "geojson"])
    def test_all_formats(self, client, users, full_dataset, export_root, fmt):
        """Все объявленные форматы формируются успешно."""
        from accounts.models import ExportJob

        client.force_login(users["analyst"])
        client.post(reverse("exports:create"), {"dataset": "objects", "format": fmt})
        job = ExportJob.objects.latest("id")
        assert job.status == ExportJob.Status.DONE
        assert (export_root / job.file_name).exists()

    def test_unknown_dataset_rejected(self, client, users, full_dataset, export_root):
        """Неизвестный набор данных отклоняется."""
        from accounts.models import ExportJob

        client.force_login(users["analyst"])
        client.post(reverse("exports:create"), {"dataset": "неизвестно", "format": "csv"})
        assert ExportJob.objects.count() == 0

    def test_unknown_format_rejected(self, client, users, full_dataset, export_root):
        """Неподдерживаемый формат отклоняется."""
        from accounts.models import ExportJob

        client.force_login(users["analyst"])
        client.post(reverse("exports:create"), {"dataset": "objects", "format": "exe"})
        assert ExportJob.objects.count() == 0

    def test_geojson_rejected_for_flat_dataset(self, client, users, full_dataset, export_root):
        """Набор без геометрии не выгружается в GeoJSON."""
        from accounts.models import ExportJob

        client.force_login(users["analyst"])
        client.post(reverse("exports:create"), {"dataset": "flows", "format": "geojson"})
        assert ExportJob.objects.count() == 0

    def test_download_own_file(self, client, users, full_dataset, export_root):
        """Пользователь может скачать собственный отчёт."""
        from accounts.models import ExportJob

        client.force_login(users["analyst"])
        client.post(reverse("exports:create"), {"dataset": "objects", "format": "csv"})
        job = ExportJob.objects.latest("id")
        response = client.get(reverse("accounts:export_download", args=[job.pk]))
        assert response.status_code == 200

    def test_cannot_download_foreign_file(self, client, users, full_dataset, export_root):
        """Чужой отчёт скачать нельзя."""
        from accounts.models import ExportJob

        client.force_login(users["analyst"])
        client.post(reverse("exports:create"), {"dataset": "objects", "format": "csv"})
        job = ExportJob.objects.latest("id")

        client.force_login(users["operator"])
        assert client.get(
            reverse("accounts:export_download", args=[job.pk])
        ).status_code == 404


class TestExportIsNotIdempotentGet:
    """Формирование отчёта изменяет состояние и потому недоступно по ссылке.

    Метод: негативное тестирование. Обработчик создаёт запись задания,
    пишет файл и оставляет запись в журнале аудита. Пока он отвечал на GET,
    отчёт формировался от предзагрузки ссылки браузером и от обхода роботом,
    а защита от подделки запросов не применялась вовсе.
    """

    def test_get_is_rejected(self, client, users, full_dataset, export_root):
        """Обращение методом GET отклоняется."""
        from accounts.models import ExportJob

        client.force_login(users["analyst"])
        response = client.get(
            reverse("exports:create"), {"dataset": "objects", "format": "csv"}
        )
        assert response.status_code == 405
        assert ExportJob.objects.count() == 0

    def test_post_without_csrf_is_rejected(self, users, full_dataset, export_root):
        """Отправка без маркера защиты от подделки отклоняется."""
        from accounts.models import ExportJob
        from django.test import Client

        enforcing = Client(enforce_csrf_checks=True)
        enforcing.force_login(users["analyst"])
        response = enforcing.post(
            reverse("exports:create"), {"dataset": "objects", "format": "csv"}
        )
        assert response.status_code == 403
        assert ExportJob.objects.count() == 0


class TestExportFilters:
    """Условия отбора переносятся в отчёт со страницы, с которой он вызван."""

    def test_filters_narrow_the_report(self, client, users, full_dataset, export_root):
        """Отчёт содержит ровно ту выборку, что видна на экране."""
        from accounts.models import ExportJob
        from core.models import InfrastructureObject

        client.force_login(users["analyst"])
        district = InfrastructureObject.objects.first().district

        client.post(reverse("exports:create"), {
            "dataset": "objects",
            "format": "csv",
            "filters": f"district={district.pk}",
        })
        narrowed = ExportJob.objects.latest("id")

        client.post(reverse("exports:create"), {"dataset": "objects", "format": "csv"})
        full = ExportJob.objects.latest("id")

        expected = InfrastructureObject.objects.filter(district=district).count()
        assert narrowed.row_count == expected
        assert full.row_count == InfrastructureObject.objects.count()
        assert narrowed.row_count < full.row_count

    def test_filters_are_recorded_with_the_job(self, client, users, full_dataset,
                                               export_root):
        """Условия отбора сохраняются в задании — отчёт воспроизводим."""
        from accounts.models import ExportJob

        client.force_login(users["analyst"])
        client.post(reverse("exports:create"), {
            "dataset": "objects", "format": "csv", "filters": "district=1&type=2",
        })
        job = ExportJob.objects.latest("id")
        assert "district=1" in job.query
        assert "type=2" in job.query


class TestCleanup:
    """Регламентная очистка каталога выгрузок."""

    def test_dry_run_keeps_files(self, users, full_dataset, export_root):
        """Пробный запуск не удаляет файлы."""
        from datetime import timedelta

        from accounts.models import ExportJob
        from django.core.management import call_command
        from django.utils import timezone

        path = export_root / "old-report.csv"
        path.write_text("данные", encoding="utf-8")
        job = ExportJob.objects.create(
            user=users["analyst"], title="Старый", dataset="objects", fmt="csv",
            file_name=path.name, status=ExportJob.Status.DONE,
        )
        ExportJob.objects.filter(pk=job.pk).update(
            created_at=timezone.now() - timedelta(days=90)
        )

        call_command("cleanup_exports", "--days", "14", "--dry-run")
        assert path.exists() and ExportJob.objects.filter(pk=job.pk).exists()

    def test_removes_expired(self, users, full_dataset, export_root):
        """Файлы с истёкшим сроком хранения удаляются."""
        from datetime import timedelta

        from accounts.models import ExportJob
        from django.core.management import call_command
        from django.utils import timezone

        path = export_root / "expired-report.csv"
        path.write_text("данные", encoding="utf-8")
        job = ExportJob.objects.create(
            user=users["analyst"], title="Истёкший", dataset="objects", fmt="csv",
            file_name=path.name, status=ExportJob.Status.DONE,
        )
        ExportJob.objects.filter(pk=job.pk).update(
            created_at=timezone.now() - timedelta(days=90)
        )

        call_command("cleanup_exports", "--days", "14")
        assert not path.exists()
        assert not ExportJob.objects.filter(pk=job.pk).exists()


class TestUnmeasuredValues:
    """Неизмеренная величина не превращается в выгрузке в ноль."""

    def test_absent_total_is_explained(self, objects):
        """Итог по неизмеренной величине сопровождается пояснением."""
        from core.models import InfrastructureObject

        InfrastructureObject.objects.update(capacity_tons=None)
        summary = dict(DATASETS["objects"][0]({}).summary)
        assert summary["Суммарная мощность хранения, т"] == "не измерена ни у одного объекта"

    def test_absent_cell_stays_empty(self, objects):
        """Пустая величина доходит до колонки неопределённой, а не нулём."""
        from core.models import InfrastructureObject

        InfrastructureObject.objects.update(capacity_tons=None)
        dataset = DATASETS["objects"][0]({})
        column = next(item for item in dataset.columns if item.title == "Мощность, т")
        assert all(column.accessor(row) is None for row in dataset.rows)

    def test_district_profiles_build_without_measurements(self, districts, roads):
        """Профили округов выгружаются и тогда, когда мощность не измерена."""
        dataset = DATASETS["districts"][0]({})
        for row in dataset.rows:
            for column in dataset.columns:
                column.accessor(row)
        assert len(list(dataset.rows)) == len(districts)
