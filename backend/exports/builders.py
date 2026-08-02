"""Формирование отчётных документов.

Модуль отвечает за выгрузку данных системы в пять форматов:

* **XLSX** — электронная таблица с форматированием, закреплённой шапкой и
  автоподбором ширины колонок; основной формат для дальнейшей обработки;
* **DOCX** — текстовый документ с титульной частью, сводкой и таблицей;
  предназначен для включения в отчётные материалы;
* **PDF** — документ фиксированной разметки для печати и рассылки;
* **CSV** — простейший формат обмена, пригодный для импорта в любую систему;
* **GeoJSON** — пространственный слой для картографических приложений.

Общая схема работы одинакова для всех форматов: набор данных описывается
объектом :class:`Dataset`, содержащим заголовки колонок, функцию извлечения
значений и выборку. Конкретный построитель отвечает только за представление,
но не за состав данных, — благодаря этому добавление формата не затрагивает
логику отбора, а добавление набора данных не затрагивает построители.

Файлы сохраняются в каталог ``EXPORT_ROOT`` и выдаются пользователю через
представление личного кабинета, которое проверяет принадлежность файла.
Прямой доступ к каталогу по HTTP не предоставляется.
"""

from __future__ import annotations

import csv
import json
import re
import unicodedata
from collections.abc import Callable, Iterable
from dataclasses import dataclass, field
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from typing import Any

from django.conf import settings
from django.utils import timezone

# Ширина колонки в электронной таблице ограничивается сверху: длинные адреса
# иначе растягивают лист до нечитаемого состояния.
MAX_COLUMN_WIDTH = 52
MIN_COLUMN_WIDTH = 10


@dataclass
class Column:
    """Описание колонки отчёта.

    Атрибуты:
        title: заголовок колонки в документе;
        accessor: функция извлечения значения из записи;
        width: рекомендуемая ширина в символах;
        numeric: признак числовой колонки — влияет на выравнивание.
    """

    title: str
    accessor: Callable[[Any], Any]
    width: int = 18
    numeric: bool = False


@dataclass
class Dataset:
    """Набор данных для выгрузки.

    Атрибуты:
        code: внутренний код набора, попадающий в имя файла;
        title: наименование отчёта;
        description: пояснение, выводимое в документах DOCX и PDF;
        columns: состав колонок;
        rows: последовательность записей;
        summary: пары «показатель — значение» для сводной части документа.
    """

    code: str
    title: str
    columns: list[Column]
    rows: Iterable[Any]
    description: str = ""
    summary: list[tuple[str, Any]] = field(default_factory=list)

    def materialize(self) -> list[list[Any]]:
        """Развернуть выборку в список строк со значениями колонок."""
        return [[column.accessor(row) for column in self.columns] for row in self.rows]


# ---------------------------------------------------------------------------
#  Служебные функции
# ---------------------------------------------------------------------------


def _slugify(text: str) -> str:
    """Привести строку к безопасному для файловой системы виду.

    Кириллица транслитерируется, пробелы и знаки препинания заменяются
    дефисом. Собственная реализация выбрана вместо ``django.utils.text.slugify``
    потому, что последняя удаляет кириллицу целиком, оставляя пустую строку.
    """
    table = {
        "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "e",
        "ж": "zh", "з": "z", "и": "i", "й": "y", "к": "k", "л": "l", "м": "m",
        "н": "n", "о": "o", "п": "p", "р": "r", "с": "s", "т": "t", "у": "u",
        "ф": "f", "х": "h", "ц": "c", "ч": "ch", "ш": "sh", "щ": "sch",
        "ъ": "", "ы": "y", "ь": "", "э": "e", "ю": "yu", "я": "ya",
    }
    lowered = unicodedata.normalize("NFKC", text).lower()
    converted = "".join(table.get(char, char) for char in lowered)
    cleaned = re.sub(r"[^a-z0-9]+", "-", converted).strip("-")
    return cleaned[:60] or "report"


def build_filename(dataset: Dataset, fmt: str) -> str:
    """Сформировать имя файла отчёта с отметкой времени.

    Отметка времени в имени исключает совпадение имён при повторных выгрузках
    одного набора и позволяет упорядочить файлы в каталоге хронологически.
    """
    stamp = timezone.localtime().strftime("%Y%m%d-%H%M%S")
    return f"freightflow-{_slugify(dataset.code)}-{stamp}.{fmt}"


def ensure_export_root() -> Path:
    """Создать каталог выгрузок, если он отсутствует."""
    root = Path(settings.EXPORT_ROOT)
    root.mkdir(parents=True, exist_ok=True)
    return root


def _to_number(value: Any) -> float | int | str | None:
    """Привести значение к числу, пригодному для записи в ячейку таблицы."""
    if value is None or value == "":
        return None
    if isinstance(value, Decimal):
        return float(value)
    if isinstance(value, (int, float)):
        return value
    return value


def _to_text(value: Any) -> str:
    """Привести значение к строке для текстовых форматов."""
    if value is None:
        return ""
    if isinstance(value, datetime):
        return timezone.localtime(value).strftime("%d.%m.%Y %H:%M")
    if isinstance(value, Decimal):
        return f"{value:.2f}".rstrip("0").rstrip(".")
    if isinstance(value, bool):
        return "да" if value else "нет"
    return str(value)


# ---------------------------------------------------------------------------
#  Построители форматов
# ---------------------------------------------------------------------------


def build_xlsx(dataset: Dataset, path: Path) -> int:
    """Сформировать электронную таблицу XLSX.

    Лист оформляется в соответствии с оформлением системы: тёмная шапка,
    янтарная линия под ней, закреплённая первая строка и включённый
    автофильтр. Возвращает число записанных строк данных.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = dataset.title[:31] or "Отчёт"

    header_fill = PatternFill("solid", fgColor="1C242C")
    header_font = Font(color="F2A03D", bold=True, size=10)
    accent_border = Border(bottom=Side(style="medium", color="F2A03D"))

    for index, column in enumerate(dataset.columns, start=1):
        cell = sheet.cell(row=1, column=index, value=column.title)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = accent_border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        sheet.column_dimensions[get_column_letter(index)].width = max(
            MIN_COLUMN_WIDTH, min(column.width, MAX_COLUMN_WIDTH)
        )

    rows = dataset.materialize()
    for row_index, values in enumerate(rows, start=2):
        for column_index, (column, value) in enumerate(zip(dataset.columns, values, strict=False), start=1):
            cell = sheet.cell(row=row_index, column=column_index)
            if column.numeric:
                cell.value = _to_number(value)
                cell.alignment = Alignment(horizontal="right")
                cell.number_format = "# ##0.00" if isinstance(cell.value, float) else "# ##0"
            else:
                cell.value = _to_text(value)
                cell.alignment = Alignment(vertical="top", wrap_text=True)

    sheet.freeze_panes = "A2"
    if rows:
        sheet.auto_filter.ref = (
            f"A1:{get_column_letter(len(dataset.columns))}{len(rows) + 1}"
        )
    sheet.row_dimensions[1].height = 30

    # Второй лист содержит сведения о выгрузке: состав отбора, время и автора.
    info = workbook.create_sheet("Сведения о выгрузке")
    info.column_dimensions["A"].width = 34
    info.column_dimensions["B"].width = 60
    meta = [
        ("Наименование отчёта", dataset.title),
        ("Пояснение", dataset.description),
        ("Число записей", len(rows)),
        ("Сформирован", timezone.localtime().strftime("%d.%m.%Y %H:%M")),
        ("Информационная система", f"{settings.PROJECT_NAME} ({settings.PROJECT_NAME_LATIN})"),
        ("Разработчик", settings.PROJECT_AUTHOR),
    ] + [(title, _to_text(value)) for title, value in dataset.summary]

    for row_index, (title, value) in enumerate(meta, start=1):
        info.cell(row=row_index, column=1, value=title).font = Font(bold=True, size=10)
        info.cell(row=row_index, column=2, value=_to_text(value)).alignment = Alignment(
            wrap_text=True, vertical="top"
        )

    workbook.save(path)
    return len(rows)


def build_docx(dataset: Dataset, path: Path) -> int:
    """Сформировать текстовый документ DOCX.

    Документ содержит заголовок, пояснение, сводку показателей и таблицу
    данных. Разметка альбомная: отчёты системы, как правило, широкие.
    """
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    document = Document()

    # Альбомная ориентация: ширина таблиц отчётов превышает книжный формат.
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.5)

    heading = document.add_heading(dataset.title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = document.add_paragraph()
    run = subtitle.add_run(
        f"Информационная система «{settings.PROJECT_NAME}» · "
        f"сформировано {timezone.localtime().strftime('%d.%m.%Y в %H:%M')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x78, 0x85)

    if dataset.description:
        document.add_paragraph(dataset.description)

    rows = dataset.materialize()

    if dataset.summary:
        document.add_heading("Сводные показатели", level=2)
        summary_table = document.add_table(rows=0, cols=2)
        summary_table.style = "Light List Accent 1"
        for title, value in dataset.summary:
            cells = summary_table.add_row().cells
            cells[0].text = title
            cells[1].text = _to_text(value)

    document.add_heading("Данные", level=2)
    table = document.add_table(rows=1, cols=len(dataset.columns))
    table.style = "Light Grid Accent 1"

    for index, column in enumerate(dataset.columns):
        cell = table.rows[0].cells[index]
        cell.text = column.title
        for paragraph in cell.paragraphs:
            for cell_run in paragraph.runs:
                cell_run.font.bold = True
                cell_run.font.size = Pt(8)

    for values in rows:
        cells = table.add_row().cells
        for index, (column, value) in enumerate(zip(dataset.columns, values, strict=False)):
            cells[index].text = _to_text(value)
            for paragraph in cells[index].paragraphs:
                if column.numeric:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for cell_run in paragraph.runs:
                    cell_run.font.size = Pt(8)

    footer = document.add_paragraph()
    footer_run = footer.add_run(
        f"Всего записей: {len(rows)}. Разработчик системы: {settings.PROJECT_AUTHOR}."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x6B, 0x78, 0x85)

    document.save(path)
    return len(rows)


def build_pdf(dataset: Dataset, path: Path) -> int:
    """Сформировать документ PDF.

    Используется встроенный шрифт с поддержкой кириллицы; при его отсутствии
    в системе применяется запасной вариант, поэтому формирование документа
    не прерывается.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    font_name = "Helvetica"
    for candidate in (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ):
        if Path(candidate).exists():
            try:
                pdfmetrics.registerFont(TTFont("Report", candidate))
                font_name = "Report"
                break
            except Exception:  # pragma: no cover — повреждённый файл шрифта
                continue

    document = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title=dataset.title,
        author=settings.PROJECT_AUTHOR,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle", parent=styles["Heading1"], fontName=font_name, fontSize=15
    )
    text_style = ParagraphStyle(
        "ReportText", parent=styles["Normal"], fontName=font_name, fontSize=8, leading=10
    )
    cell_style = ParagraphStyle(
        "ReportCell", parent=text_style, fontSize=7, leading=8.5
    )

    story = [
        Paragraph(dataset.title, title_style),
        Paragraph(
            f"Информационная система «{settings.PROJECT_NAME}» · сформировано "
            f"{timezone.localtime().strftime('%d.%m.%Y в %H:%M')}",
            text_style,
        ),
        Spacer(1, 6 * mm),
    ]

    if dataset.description:
        story.extend([Paragraph(dataset.description, text_style), Spacer(1, 4 * mm)])

    rows = dataset.materialize()

    # Число строк в PDF ограничивается: документ на несколько тысяч страниц
    # практической ценности не имеет, для полной выгрузки служит XLSX или CSV.
    limit = 400
    visible = rows[:limit]

    data = [[Paragraph(column.title, cell_style) for column in dataset.columns]]
    data += [
        [Paragraph(_to_text(value), cell_style) for value in values] for values in visible
    ]

    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1C242C")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#F2A03D")),
                ("LINEBELOW", (0, 0), (-1, 0), 1, colors.HexColor("#F2A03D")),
                ("GRID", (0, 1), (-1, -1), 0.25, colors.HexColor("#D5DCDF")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#F7F9F9")]),
            ]
        )
    )
    story.append(table)

    if len(rows) > limit:
        story.extend([
            Spacer(1, 4 * mm),
            Paragraph(
                f"Показаны первые {limit} записей из {len(rows)}. "
                "Полная выборка доступна в форматах XLSX и CSV.",
                text_style,
            ),
        ])

    story.extend([
        Spacer(1, 6 * mm),
        Paragraph(f"Разработчик системы: {settings.PROJECT_AUTHOR}", text_style),
    ])

    document.build(story)
    return len(rows)


def build_csv(dataset: Dataset, path: Path) -> int:
    """Сформировать таблицу CSV.

    Файл записывается в кодировке UTF-8 с меткой порядка байтов и разделителем
    «точка с запятой»: в таком виде он открывается в русскоязычной версии
    табличного процессора без дополнительных настроек импорта.
    """
    rows = dataset.materialize()
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=";", quoting=csv.QUOTE_MINIMAL)
        writer.writerow([column.title for column in dataset.columns])
        for values in rows:
            writer.writerow([_to_text(value) for value in values])
    return len(rows)


def build_geojson(dataset: Dataset, path: Path, geometry_getter: Callable) -> int:
    """Сформировать пространственный слой GeoJSON.

    Записи без геометрии в слой не включаются: объект без координат не может
    быть отображён на карте, а его присутствие в файле нарушило бы структуру
    коллекции.
    """
    from geo import Geometry

    features = []
    for row in dataset.rows:
        geometry = geometry_getter(row)
        if not isinstance(geometry, Geometry):
            continue
        properties = {
            column.title: _to_text(column.accessor(row)) for column in dataset.columns
        }
        features.append(geometry.as_feature(properties))

    payload = {
        "type": "FeatureCollection",
        "name": dataset.title,
        "crs": {"type": "name", "properties": {"name": "urn:ogc:def:crs:OGC:1.3:CRS84"}},
        "metadata": {
            "system": settings.PROJECT_NAME,
            "generated_at": timezone.localtime().isoformat(),
            "author": settings.PROJECT_AUTHOR,
            "count": len(features),
        },
        "features": features,
    }

    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=1), encoding="utf-8"
    )
    return len(features)


#: Соответствие кода формата и построителя документа.
BUILDERS: dict[str, Callable] = {
    "xlsx": build_xlsx,
    "docx": build_docx,
    "pdf": build_pdf,
    "csv": build_csv,
}
